# imports
from docx import Document
from helper import write_text_units
import logging

# logging configuration
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

document = Document('./data/code_of_studies.docx')

# The maximum number of characters a chunk can contain
max_chunks_size = 2000 

# The 111st paragraph is the first chapter
first_par = 111 

last_par = len(document.paragraphs)

def slice_document():
    logging.info("Slicing document to chunks started...")
    chunks = [] # list of all the chunks
    chunk = ""
    idx = first_par
    while idx < last_par:
        curr_par = document.paragraphs[idx]
        if curr_par.style.name == 'Heading 2': # section changed!
            chunks.append(chunk)
            chunk = ""
        if len(chunk) >= max_chunks_size:
            chunks.append(chunk)
            chunk = ""
        text = curr_par.text.replace("\n", "").replace("\r", "") # removing line breaks
        chunk += f'{text} '
        idx += 1
    return chunks

def create_chunks():
    logging.info("Creating the chunks started...")
    chunks = slice_document()
    write_text_units("chunks", chunks)